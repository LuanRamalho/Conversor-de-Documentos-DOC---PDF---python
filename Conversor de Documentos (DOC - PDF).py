import os
import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter.ttk import Combobox
from docx import Document
import pdfplumber
from fpdf import FPDF

# Funções para manipulação de arquivos
def doc_to_pdf(input_path, output_path):
    try:
        document = Document(input_path)
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for paragraph in document.paragraphs:
            pdf.multi_cell(0, 10, paragraph.text)
        pdf.output(output_path)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao converter DOCX para PDF: {e}")

def pdf_to_doc(input_path, output_path):
    try:
        document = Document()
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                document.add_paragraph(text)
        document.save(output_path)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao converter PDF para DOCX: {e}")

def convert_file(input_path, input_format, output_format):
    base_name = os.path.splitext(input_path)[0]
    output_path = f"{base_name}_converted.{output_format.lower()}"

    if input_format == "DOCX" and output_format == "PDF":
        doc_to_pdf(input_path, output_path)
    elif input_format == "PDF" and output_format == "DOCX":
        pdf_to_doc(input_path, output_path)
    else:
        messagebox.showerror("Erro", "Conversão não suportada no momento.")
        return

    messagebox.showinfo("Sucesso", f"Arquivo convertido e salvo em: {output_path}")

# Função para selecionar arquivo
def select_file():
    file_path = filedialog.askopenfilename(
        filetypes=[("Arquivos suportados", "*.docx;*.pdf"), ("Todos os arquivos", "*.*")]
    )
    if file_path:
        input_file_entry.delete(0, tk.END)
        input_file_entry.insert(0, file_path)

# Função para iniciar conversão
def start_conversion():
    input_path = input_file_entry.get()
    input_format = input_format_combo.get()
    output_format = output_format_combo.get()

    if not os.path.isfile(input_path):
        messagebox.showerror("Erro", "Por favor, selecione um arquivo válido.")
        return

    if input_format == output_format:
        messagebox.showerror("Erro", "Formato de entrada e saída não podem ser iguais.")
        return

    convert_file(input_path, input_format, output_format)

# Configuração da interface gráfica
root = tk.Tk()
root.title("Conversor de Arquivos")
root.geometry("400x300")
root.configure(bg="#71FFC8")

# Título
title_label = tk.Label(root, text="Conversor de Arquivos", font=("Helvetica", 16), bg="#71FFC8", fg="#333")
title_label.pack(pady=10)

# Seção de seleção de arquivo
frame_file = tk.Frame(root, bg="#71FFC8")
frame_file.pack(pady=5)
file_label = tk.Label(frame_file, text="Arquivo:", bg="#71FFC8")
file_label.pack(side=tk.LEFT)
input_file_entry = tk.Entry(frame_file, width=30)
input_file_entry.pack(side=tk.LEFT, padx=5)
select_file_button = tk.Button(frame_file, text="Selecionar", command=select_file, bg="#4CAF50", fg="white")
select_file_button.pack(side=tk.LEFT)

# Combobox para formatos de entrada e saída
frame_format = tk.Frame(root, bg="#71FFC8")
frame_format.pack(pady=10)

input_format_label = tk.Label(frame_format, text="Formato de Entrada:", bg="#71FFC8")
input_format_label.grid(row=0, column=0, padx=5)
input_format_combo = Combobox(frame_format, values=["DOCX", "PDF"], state="readonly")
input_format_combo.grid(row=0, column=1, padx=5)
input_format_combo.current(0)

output_format_label = tk.Label(frame_format, text="Formato de Saída:", bg="#71FFC8")
output_format_label.grid(row=1, column=0, padx=5)
output_format_combo = Combobox(frame_format, values=["DOCX", "PDF"], state="readonly")
output_format_combo.grid(row=1, column=1, padx=5)
output_format_combo.current(1)

# Botão para realizar a conversão
convert_button = tk.Button(root, text="Converter", command=start_conversion, bg="#2196F3", fg="white", font=("Helvetica", 12))
convert_button.pack(pady=20)

# Iniciar a interface
def main():
    root.mainloop()

if __name__ == "__main__":
    main()